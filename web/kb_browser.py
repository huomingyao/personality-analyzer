"""知识库浏览器 - Agent 主动搜索和浏览知识库的接口"""

from __future__ import annotations

import os
import json
from pathlib import Path
from dataclasses import dataclass, field
from typing import Any, List
from collections import defaultdict

# 路径常量
KB_BASE = Path(__file__).parent / "knowledge_bases"


@dataclass
class KbDocument:
    """知识库文档"""
    doc_id: str
    title: str
    content: str
    file_path: str
    sections: List[str] = field(default_factory=list)


@dataclass
class KbSearchResult:
    """知识库搜索结果"""
    query: str
    documents: List[KbDocument] = field(default_factory=list)
    total_matches: int = 0
    search_method: str = ""  # "keyword" / "semantic" / "tree"


class KnowledgeBrowser:
    """Agent 使用的知识库浏览器 - 主动搜索知识库内容"""

    def __init__(self, kb_name: str | None = None):
        self.kb_name = kb_name
        self._kb_tree: dict | None = None
        self._documents: dict[str, KbDocument] = {}

    @property
    def kb_path(self) -> Path | None:
        if not self.kb_name:
            return None
        return KB_BASE / self.kb_name

    # ==================== 列举知识库 ====================

    def list_knowledge_bases(self) -> list[dict]:
        """列出所有可用知识库"""
        if not KB_BASE.exists():
            return []

        bases = []
        for name in KB_BASE.iterdir():
            if name.is_dir():
                files = [f.name for f in name.iterdir() if f.is_file()]
                bases.append({
                    "name": name.name,
                    "file_count": len(files),
                })
        return bases

    def list_documents(self) -> list[str]:
        """列出现有知识库中的所有文档"""
        if not self.kb_path or not self.kb_path.exists():
            return []

        docs = []
        for f in self.kb_path.iterdir():
            if f.is_file() and not f.name.startswith('.'):
                docs.append(f.stem)
        return docs

    # ==================== 浏览知识树 ====================

    def load_kb_tree(self) -> dict:
        """加载知识库的索引树（Corpus2Skill 生成的）"""
        if not self.kb_path:
            return {}

        pipeline_dir = self.kb_path / ".pipeline"
        if not pipeline_dir.exists():
            return {}

        index_file = pipeline_dir / "INDEX.md"
        if index_file.exists():
            tree_content = index_file.read_text(encoding='utf-8')
            return {
                "kb_name": self.kb_name,
                "index_md": tree_content,
            }

        return {"kb_name": self.kb_name}

    def get_kb_structure(self) -> dict:
        """获取知识库的完整结构"""
        if not self.kb_path or not self.kb_path.exists():
            return {"error": f"知识库 [{self.kb_name}] 不存在"}

        docs = []
        pipeline_dir = self.kb_path / ".pipeline"

        # 读取已处理的 chunks
        if pipeline_dir.exists():
            for chunk_file in pipeline_dir.glob("*_chunks.json"):
                doc_id = chunk_file.stem.replace("_chunks", "")
                data = json.loads(chunk_file.read_text(encoding='utf-8'))

                # 获取章节标题列表
                sections = list(set(
                    c.get("section_title", "")
                    for c in data
                    if c.get("section_title")
                ))

                docs.append({
                    "doc_id": doc_id,
                    "chunk_count": len(data),
                    "sections": sections[:10],  # 限制数量
                })

        return {
            "kb_name": self.kb_name,
            "documents": docs,
        }

    # ==================== 读取文档内容 ====================

    def read_document(self, doc_id: str) -> KbDocument | None:
        """读取指定文档的完整内容"""
        if not self.kb_path:
            return None

        # 首先尝试 Pipeline 处理过的 chunks
        pipeline_dir = self.kb_path / ".pipeline"
        chunks_file = pipeline_dir / f"{doc_id}_chunks.json"

        if chunks_file.exists():
            chunks_data = json.loads(chunks_file.read_text(encoding='utf-8'))
            content = "\n\n".join(c["content"] for c in chunks_data)
            sections = list(set(
                c.get("section_title", "")
                for c in chunks_data
                if c.get("section_title")
            ))
            return KbDocument(
                doc_id=doc_id,
                title=doc_id,
                content=content,
                file_path=str(chunks_file),
                sections=sections,
            )

        # 回退：读取原始文件
        for f in self.kb_path.iterdir():
            if f.is_file() and f.stem == doc_id:
                ext = f.suffix.lower()
                content = self._extract_text(f)
                return KbDocument(
                    doc_id=doc_id,
                    title=doc_id,
                    content=content,
                    file_path=str(f),
                )

        return None

    def read_document_by_name(self, filename: str) -> str:
        """按文件名读取内容"""
        if not self.kb_path:
            return ""

        file_path = self.kb_path / filename
        if file_path.exists():
            return self._extract_text(file_path)

        # 尝试多种扩展名
        for ext in ['.txt', '.md', '.docx', '.pdf']:
            file_path = self.kb_path / f"{filename}{ext}"
            if file_path.exists():
                return self._extract_text(file_path)

        return ""

    # ==================== 关键词搜索 ====================

    def search_by_keyword(self, query: str, max_docs: int = 5) -> KbSearchResult:
        """按关键词搜索知识库"""
        if not self.kb_path:
            return KbSearchResult(query=query, search_method="keyword")

        query_lower = query.lower()
        matched_docs: list[KbDocument] = []

        # 搜索 Pipeline chunks
        pipeline_dir = self.kb_path / ".pipeline"

        if pipeline_dir.exists():
            for chunk_file in pipeline_dir.glob("*_chunks.json"):
                doc_id = chunk_file.stem.replace("_chunks", "")
                data = json.loads(chunk_file.read_text(encoding='utf-8'))

                # 在 chunks 中搜索关键词
                matched_chunks = [
                    c for c in data
                    if query_lower in c.get("content", "").lower()
                ]

                if matched_chunks:
                    # 合并匹配的 chunks
                    combined = "\n\n".join(c["content"] for c in matched_chunks[:10])
                    sections = list(set(
                        c.get("section_title", "")
                        for c in matched_chunks
                        if c.get("section_title")
                    ))

                    matched_docs.append(KbDocument(
                        doc_id=doc_id,
                        title=doc_id,
                        content=combined,
                        file_path=str(chunk_file),
                        sections=sections,
                    ))

                    if len(matched_docs) >= max_docs:
                        break

        return KbSearchResult(
            query=query,
            documents=matched_docs,
            total_matches=len(matched_docs),
            search_method="keyword",
        )

    def search_all_keywords(self, queries: list[str], max_per_query: int = 3) -> dict[str, KbSearchResult]:
        """一次搜索多个关键词"""
        results = {}
        for q in queries:
            results[q] = self.search_by_keyword(q, max_per_query)
        return results

    # ==================== 全文检索 ====================

    def search_full_text(self, query: str) -> list[dict]:
        """全文检索，返回匹配的片段"""
        if not self.kb_path:
            return []

        matches = []
        query_lower = query.lower()

        pipeline_dir = self.kb_path / ".pipeline"

        if pipeline_dir.exists():
            for chunk_file in pipeline_dir.glob("*_chunks.json"):
                doc_id = chunk_file.stem.replace("_chunks", "")
                data = json.loads(chunk_file.read_text(encoding='utf-8'))

                for chunk in data:
                    content = chunk.get("content", "")
                    if query_lower in content.lower():
                        # 找到上下文片段（前后各100字）
                        idx = content.lower().find(query_lower)
                        start = max(0, idx - 100)
                        end = min(len(content), idx + len(query) + 100)
                        snippet = content[start:end]

                        matches.append({
                            "doc_id": doc_id,
                            "chunk_id": chunk.get("id", ""),
                            "section_title": chunk.get("section_title", ""),
                            "snippet": f"...{snippet}...",
                        })

        return matches

    # ==================== 辅助方法 ====================

    def _extract_text(self, file_path: Path) -> str:
        """提取文本内容"""
        ext = file_path.suffix.lower()
        try:
            if ext == '.docx':
                from docx import Document
                doc = Document(str(file_path))
                return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            elif ext == '.pdf':
                import fitz
                text = []
                with fitz.open(str(file_path)) as pdf:
                    for page in pdf:
                        text.append(page.get_text())
                return '\n'.join(text)
            else:
                # 文本文件
                for enc in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                    try:
                        return file_path.read_text(encoding=enc)
                    except UnicodeDecodeError:
                        continue
                return file_path.read_bytes().decode('utf-8', errors='ignore')
        except Exception:
            return "[无法解析的文件]"

    def get_summary(self) -> str:
        """获取知识库摘要信息"""
        bases = self.list_knowledge_bases()
        if self.kb_name:
            docs = self.list_documents()
            structure = self.get_kb_structure()
            doc_count = len(structure.get("documents", []))
            return f"知识库 [{self.kb_name}] - {doc_count} 篇文档"
        else:
            return f"可用知识库: {', '.join(b['name'] for b in bases)}"


# ==================== 便捷工厂函数 ====================

def create_browser(kb_name: str | None = None) -> KnowledgeBrowser:
    """创建知识库浏览器实例"""
    return KnowledgeBrowser(kb_name)


def search_kb(kb_name: str, query: str, method: str = "keyword") -> KbSearchResult:
    """快速搜索函数"""
    browser = KnowledgeBrowser(kb_name)
    return browser.search_by_keyword(query)